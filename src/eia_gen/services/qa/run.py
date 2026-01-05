from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Any

from pydantic import BaseModel

from eia_gen.models.case import Case, ScopingClass
from eia_gen.models.fields import QuantityField
from eia_gen.models.fields import TextField
from eia_gen.models.sources import SourceRegistry
from eia_gen.services.figures.spec_figures import build_figure_map, is_required
from eia_gen.services.figures.materialize import select_pdf_page
from eia_gen.services.qa.report import RuleResult, ValidationReport
from eia_gen.services.tables.path import resolve_path
from eia_gen.services.tables.spec_tables import build_table
from eia_gen.services.tables.validations import ValidationFinding, validate_table
from eia_gen.spec.models import SpecBundle


_CITATION_BLOCK_RE = re.compile(r"〔([^〕]+)〕")
_AUTH_REFERENCE_RE = re.compile(r"(?i)AUTHENTICITY\s*[:=]\s*REFERENCE")

_FIELD_CLAIM_PATTERNS: list[tuple[str, re.Pattern[str], str]] = [
    ("E-FIELD-001", re.compile(r"(현지|현장)조사\s*결과"), "현장/현지조사 결과 표현"),
    ("E-FIELD-002", re.compile(r"(현지|현장)조사[^.\n]{0,30}(실시|수행)(하였|했다|함)"), "현장/현지조사 실시(과거형)"),
    ("E-FIELD-003", re.compile(r"(현장|현지)측정\s*결과"), "현장측정 결과 표현"),
    ("E-FIELD-004", re.compile(r"(탐문|인터뷰)\s*조사\s*결과"), "탐문/인터뷰 조사 결과 표현"),
    ("E-FIELD-005", re.compile(r"(탐문|인터뷰)\s*조사[^.\n]{0,30}(실시|수행)(하였|했다|함)"), "탐문/인터뷰 조사 실시(과거형)"),
]

_A3_LONG_MM = 420.0
_A3_SHORT_MM = 297.0


_EMPTY_CELL_SUBSTRINGS = (
    "【작성자 기입 필요】",
    "[작성자 기입 필요]",
    "【자료 확인 필요】",
    "[자료 확인 필요]",
)
_EMPTY_CELL_EXACT_UPPER = {"NA", "N/A", "UNKNOWN", "UNK", "TBD"}


def _is_empty_cell(v: Any) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        text = v.strip()
        if not text:
            return True
        if text.upper() in _EMPTY_CELL_EXACT_UPPER:
            return True
        if any(token in text for token in _EMPTY_CELL_SUBSTRINGS):
            return True
    return False


def _is_empty_row(row: list[Any]) -> bool:
    return all(_is_empty_cell(v) for v in row)


def _default_data_acquisition_rules_path() -> Path | None:
    """Best-effort locate `config/data_acquisition_rules.yaml` when running from repo."""
    try:
        here = Path(__file__).resolve()
    except Exception:
        return None
    for cand in [here.parent] + list(here.parents):
        p = cand / "config" / "data_acquisition_rules.yaml"
        if p.exists():
            return p
    return None


def _load_data_acquisition_rules(path: Path) -> list[dict[str, Any]]:
    try:
        import yaml  # type: ignore
    except Exception:
        return []

    try:
        obj = yaml.safe_load(path.read_text(encoding="utf-8"))
    except Exception:
        return []

    if not isinstance(obj, dict):
        return []
    rules = obj.get("rules") or []
    return [r for r in rules if isinstance(r, dict)]


def _evaluate_data_acquisition_rules(
    *, xlsx_path: Path, rules: list[dict[str, Any]], submission_mode: bool = False
) -> list[RuleResult]:
    """Evaluate `config/data_acquisition_rules.yaml` against a case.xlsx(v2).

    When `submission_mode=True`, rules may optionally provide `submission_severity`
    to tighten gating for "ready to submit" checks without blocking draft iteration.
    """
    try:
        import openpyxl  # type: ignore
    except Exception:
        return []

    try:
        wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception:
        return []

    out: list[RuleResult] = []

    def _sheet_row_count(sheet_name: str, *, any_of: list[str] | None = None) -> int:
        if sheet_name not in wb.sheetnames:
            return 0
        ws = wb[sheet_name]
        idxs: list[int] = []
        if any_of:
            headers = [c.value for c in ws[1]]
            hm = {str(h).strip(): i for i, h in enumerate(headers) if h is not None and str(h).strip()}
            idxs = [hm.get(str(c).strip()) for c in any_of if str(c).strip()]
            idxs = [i for i in idxs if isinstance(i, int)]
            # If none of the columns exist, treat as empty so the rule can still nudge the user.
            if not idxs:
                return 0
        count = 0
        for r in ws.iter_rows(min_row=2, values_only=True):
            row = list(r or [])
            if not row or _is_empty_row(row):
                continue
            if idxs:
                if not any(not _is_empty_cell(row[i] if i < len(row) else None) for i in idxs):
                    continue
            count += 1
        return count

    def _column_all_empty(sheet_name: str, column_name: str) -> tuple[bool, str]:
        """Return (is_empty, related_row_id) best-effort."""
        if sheet_name not in wb.sheetnames:
            return (True, "")
        ws = wb[sheet_name]
        headers = [c.value for c in ws[1]]
        hm = {str(h).strip(): i for i, h in enumerate(headers) if h is not None and str(h).strip()}
        idx = hm.get(column_name)
        if idx is None:
            # treat missing column as empty so rule can still nudge the user
            return (True, "")

        # Prefer the first non-empty data row for row_id linkage.
        first_row_id = ""
        for ridx, r in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            row = list(r or [])
            if not row or _is_empty_row(row):
                continue
            if not first_row_id:
                first_row_id = str(ridx)
            v = row[idx] if idx < len(row) else None
            if not _is_empty_cell(v):
                return (False, first_row_id)
        return (True, first_row_id)

    for rule in rules:
        rule_id = str(rule.get("id") or "").strip()
        when = rule.get("when") if isinstance(rule.get("when"), dict) else {}
        sheet = str(when.get("sheet") or "").strip()
        if not rule_id or not sheet:
            continue

        want_row_count = when.get("row_count")
        want_col = str(when.get("column") or "").strip()
        want_empty = bool(when.get("is_empty") is True)
        want_any_of = when.get("any_of")
        any_of_cols: list[str] | None = None
        if isinstance(want_any_of, list):
            cols = [str(c).strip() for c in want_any_of if str(c).strip()]
            any_of_cols = cols or None

        triggered = False
        related_anchor = ""
        related_row_id = ""

        if want_row_count is not None:
            try:
                n = int(want_row_count)
            except Exception:
                n = None
            if n is not None and _sheet_row_count(sheet, any_of=any_of_cols) == n:
                triggered = True
                related_anchor = sheet

        if want_col and want_empty:
            empty, row_id = _column_all_empty(sheet, want_col)
            if empty:
                triggered = True
                related_anchor = f"{sheet}.{want_col}"
                related_row_id = row_id

        if not triggered:
            continue

        severity = str(rule.get("severity") or "WARN").strip().upper()
        if submission_mode:
            sev2 = str(rule.get("submission_severity") or "").strip().upper()
            if sev2:
                severity = sev2
        if severity not in {"ERROR", "WARN", "INFO"}:
            severity = "WARN"

        msg = str(rule.get("message") or "").strip()
        if not msg:
            msg = f"입력 누락: {related_anchor}"

        suggested_sources = rule.get("suggested_sources") if isinstance(rule.get("suggested_sources"), list) else []
        hints: list[str] = []
        for it in suggested_sources:
            if not isinstance(it, dict):
                continue
            sid = str(it.get("src_id") or "").strip()
            how = str(it.get("how") or "").strip()
            if sid and how:
                hints.append(f"{sid}: {how}")
            elif sid:
                hints.append(sid)
            elif how:
                hints.append(how)

        fix = ""
        if hints:
            fix = " / ".join(hints)
        else:
            fix = "case.xlsx에서 관련 시트/컬럼을 채우고 재실행"

        out.append(
            RuleResult(
                rule_id=rule_id,
                severity=severity,  # type: ignore[arg-type]
                message=msg,
                fix_hint=fix,
                path=related_anchor or sheet,
                related_anchor=related_anchor or sheet,
                related_sheet=sheet,
                related_row_id=related_row_id,
            )
        )

    return out


def _extract_citation_ids(text: str) -> list[str]:
    ids: list[str] = []
    for m in _CITATION_BLOCK_RE.finditer(text or ""):
        inside = m.group(1)
        for raw in inside.split(","):
            token = raw.strip()
            if not token:
                continue
            if token.upper().startswith("SRC:"):
                token = token.split(":", 1)[1].strip()
            ids.append(token)
    return ids


def _has_any_citation(text: str) -> bool:
    return bool(_CITATION_BLOCK_RE.search(text or ""))


def _is_field_evidence_source_id(sources: SourceRegistry, source_id: str) -> bool:
    sid = (source_id or "").strip()
    if not sid:
        return False
    if sid in {"S-TBD", "SRC-TBD"}:
        return False

    entry = sources.get(sid)
    entry_type = (getattr(entry, "type", None) or "").strip().lower() if entry else ""
    if entry_type and any(k in entry_type for k in ("field", "survey", "photo", "interview", "현장", "탐문")):
        return True

    up = sid.upper()
    if any(token in up for token in ("SRC-FIELD", "FIELD", "SURVEY", "PHOTO", "INTERVIEW")):
        return True

    return False


def _collect_linked_source_ids_from_field_meta(case: Case) -> set[str]:
    """Best-effort collect source_ids that are actually linked in case.xlsx meta.

    Used to tighten the "field claim" guardrail: citing a field source_id should be backed
    by at least one of:
    - FIELD_SURVEY_LOG rows (field_survey_log extra)
    - ATTACHMENTS rows with a non-empty file_path (attachments_manifest extra)
    """

    out: set[str] = set()
    extra = case.model_extra or {}

    field_log = extra.get("field_survey_log")
    if isinstance(field_log, list):
        for row in field_log:
            if not isinstance(row, dict):
                continue
            for v in row.values():
                if isinstance(v, (TextField, QuantityField)):
                    for sid in (v.src or []):
                        token = str(sid).strip()
                        if token and token not in {"S-TBD", "SRC-TBD"}:
                            out.add(token)
                elif isinstance(v, dict):
                    src = v.get("src") or []
                    if isinstance(src, list):
                        for sid in src:
                            token = str(sid).strip()
                            if token and token not in {"S-TBD", "SRC-TBD"}:
                                out.add(token)

    attachments = extra.get("attachments_manifest")
    if isinstance(attachments, list):
        for row in attachments:
            if not isinstance(row, dict):
                continue
            fp = row.get("file_path")
            fp_text = ""
            if isinstance(fp, TextField):
                fp_text = str(fp.t or "")
            elif isinstance(fp, dict):
                fp_text = str(fp.get("t") or fp.get("v") or "")
            if not fp_text.strip():
                continue

            src = []
            if isinstance(fp, TextField):
                src = fp.src or []
            elif isinstance(fp, dict):
                src = fp.get("src") or []

            if isinstance(src, list):
                for sid in src:
                    token = str(sid).strip()
                    if token and token not in {"S-TBD", "SRC-TBD"}:
                        out.add(token)

    return out


def _is_reference_figure(case: Case, fig_id: str) -> bool:
    for a in getattr(case, "assets", []) or []:
        if getattr(a, "asset_id", "") != fig_id:
            continue
        auth = str(getattr(a, "authenticity", "") or "").strip().upper()
        if auth in {"REFERENCE", "REF"}:
            return True
        if auth in {"OFFICIAL"}:
            return False
        so = str(getattr(a, "source_origin", "") or "").strip().upper()
        return so in {"REFERENCE", "REF"}
    return False


def _pdf_page_size_mm(pdf_path: Path, *, page_1based: int) -> tuple[float, float] | None:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return None

    try:
        doc = fitz.open(str(pdf_path))
        try:
            idx = max(0, min(int(page_1based) - 1, doc.page_count - 1))
            page = doc.load_page(idx)
            rect = page.rect
            w_in = float(rect.width) / 72.0
            h_in = float(rect.height) / 72.0
            return (w_in * 25.4, h_in * 25.4)
        finally:
            doc.close()
    except Exception:
        return None


def _walk(obj: Any, path: str = ""):
    # Pydantic BaseModel
    if isinstance(obj, BaseModel):
        for k in obj.__class__.model_fields:
            v = getattr(obj, k)
            yield from _walk(v, f"{path}.{k}" if path else k)
        # extras
        if obj.model_extra:
            for k, v in obj.model_extra.items():
                yield from _walk(v, f"{path}.{k}" if path else k)
        return

    if isinstance(obj, QuantityField):
        yield path, obj.v, obj.src
        return

    if isinstance(obj, TextField):
        yield path, obj.t, obj.src
        return

    if isinstance(obj, dict):
        # dict numeric field style: {"v":..., "src":[...]}
        if "v" in obj:
            yield path, obj.get("v"), obj.get("src") or []
        for k, v in obj.items():
            yield from _walk(v, f"{path}.{k}" if path else str(k))
        return

    if isinstance(obj, list):
        for i, v in enumerate(obj):
            yield from _walk(v, f"{path}[{i}]")
        return


def _template_present_ids(
    *,
    case: Case,
    spec: SpecBundle,
    template_path: Path,
) -> tuple[set[str], set[str], set[str]]:
    """Return (section_ids, table_ids, figure_ids) present in the template.

    - Uses the spec's template_map anchors as the mapping universe.
    - Respects conditional section insertion (same as the DOCX renderer).
    """
    from docx import Document

    from eia_gen.services.conditions import eval_condition

    doc = Document(str(template_path))
    present_anchor_texts = {p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()}
    sections_by_id = {s.id: s for s in spec.sections.sections}

    present_sections: set[str] = set()
    present_tables: set[str] = set()
    present_figures: set[str] = set()

    for a in spec.template_map.anchors:
        anchor_text = (a.anchor or "").strip()
        if not anchor_text or anchor_text not in present_anchor_texts:
            continue

        insert = a.insert

        # Match spec_renderer.render_template_docx conditional behavior.
        if getattr(insert, "conditional", False) and getattr(insert, "type", "") == "section":
            sec_spec = sections_by_id.get(insert.id)
            if sec_spec and sec_spec.condition and not eval_condition(case, sec_spec.condition):
                continue

        if insert.type == "section":
            present_sections.add(insert.id)
        elif insert.type == "table":
            present_tables.add(insert.id)
        elif insert.type == "figure":
            present_figures.add(insert.id)

    return present_sections, present_tables, present_figures


def run_qa(
    case: Case,
    sources: SourceRegistry,
    draft,
    spec: SpecBundle | None = None,
    *,
    asset_search_dirs: list[Path] | None = None,
    template_path: Path | None = None,
    case_xlsx_path: Path | None = None,
    data_acquisition_rules_path: Path | None = None,
    submission_mode: bool = False,
) -> ValidationReport:
    results: list[RuleResult] = []

    forbidden = []
    if spec is not None:
        forbidden = list(spec.sections.doc_profile.forbidden_phrases or [])
    else:
        forbidden = ["문제 없음", "영향 없음"]

    present_section_ids: set[str] | None = None
    present_table_ids: set[str] | None = None
    present_figure_ids: set[str] | None = None

    qa_sections = list(draft.sections)
    if spec is not None and template_path is not None:
        try:
            present_section_ids, present_table_ids, present_figure_ids = _template_present_ids(
                case=case, spec=spec, template_path=template_path
            )
            qa_sections = [s for s in qa_sections if s.section_id in present_section_ids]
        except Exception:
            # Fail open: keep QA running even when template parsing fails.
            present_section_ids = None
            present_table_ids = None
            present_figure_ids = None

    all_text = "\n".join("\n".join(s.paragraphs) for s in qa_sections)
    placeholder_token = "【작성자 기입 필요】"

    # 금칙어
    for phrase in forbidden:
        if phrase and phrase in all_text:
            results.append(
                RuleResult(
                    rule_id="W-TEXT-001",
                    severity="WARN",
                    message=f"금칙어 탐지: {phrase!r}",
                    fix_hint="단정 표현을 조건부/전제 기반 표현으로 수정",
                )
            )

    # 문단 출처 태그
    for s in qa_sections:
        for i, p in enumerate(s.paragraphs):
            if not _has_any_citation(p):
                results.append(
                    RuleResult(
                        rule_id="E-CIT-001",
                        severity="ERROR",
                        message=f"출처 태그 누락(섹션 {s.section_id} 문단 {i+1})",
                        fix_hint="문단 끝에 〔SRC:ID〕(또는 〔ID〕) 추가",
                    )
                )
                break

    # 섹션별 TODO/placeholder(입력 누락) 경고
    for s in qa_sections:
        if getattr(s, "todos", None):
            for todo in s.todos:
                if not str(todo or "").strip():
                    continue
                results.append(
                    RuleResult(
                        rule_id="W-TODO-001",
                        severity="WARN",
                        path=s.section_id,
                        message=f"입력 보완 필요: {todo}",
                        fix_hint="관련 case.xlsx(v2) 입력값을 채우고 재생성",
                    )
                )

        ph_count = 0
        for p in s.paragraphs:
            ph_count += str(p or "").count(placeholder_token)
        if ph_count:
            results.append(
                RuleResult(
                    rule_id="W-PLACEHOLDER-001",
                    severity="WARN",
                    path=s.section_id,
                    message=f"{placeholder_token} 표시가 남아있습니다(섹션 {s.section_id}, {ph_count}건).",
                    fix_hint="해당 절의 INPUT 항목(case.xlsx)을 채우거나, REUSE/AUTO로 전환",
                )
            )

    # "현장조사/현지조사/현장측정 결과" 같은 표현은
    # 해당 문단의 출처(citation)가 FIELD/PHOTO/INTERVIEW 등 "현장 근거"를 포함할 때만 허용.
    field_claim_cited_src_ids: set[str] = set()
    field_claim_cited_src_ids_in_summary_or_conclusion: set[str] = set()

    def _is_summary_or_conclusion_section(section_id: str) -> bool:
        sid = (section_id or "").strip().upper()
        if sid in {"CH0_SUMMARY", "DIA0_SUMMARY"}:
            return True
        return sid.endswith("_SUMMARY") or sid.endswith("_CONCLUSION")

    for s in qa_sections:
        for i, p in enumerate(s.paragraphs):
            if not p:
                continue
            for rule_id, pat, label in _FIELD_CLAIM_PATTERNS:
                if not pat.search(p):
                    continue
                cited = _extract_citation_ids(p)
                cited_field_ids = [cid for cid in cited if _is_field_evidence_source_id(sources, cid)]
                field_claim_cited_src_ids.update(cited_field_ids)
                if cited_field_ids and _is_summary_or_conclusion_section(s.section_id):
                    field_claim_cited_src_ids_in_summary_or_conclusion.update(cited_field_ids)
                has_field = bool(cited_field_ids)
                if has_field:
                    continue
                results.append(
                    RuleResult(
                        rule_id=rule_id,
                        severity="ERROR",
                        path=f"{s.section_id}:{i+1}",
                        message=f"{label}가 탐지되었으나 현장 근거 출처가 인용되지 않았습니다.",
                        fix_hint="현장조사/측정 근거(SRC-FIELD/사진/탐문 등)를 인용하거나, 표현을 '문헌/공공DB 기반' 또는 '추가조사 필요'로 수정",
                    )
                )

    if field_claim_cited_src_ids:
        linked_src_ids = _collect_linked_source_ids_from_field_meta(case)
        missing = sorted(sid for sid in field_claim_cited_src_ids if sid not in linked_src_ids)
        if missing:
            results.append(
                RuleResult(
                    rule_id="W-FIELD-META-001",
                    severity="WARN",
                    message=(
                        "현장 근거 출처를 인용했으나(case 본문), "
                        "case.xlsx의 FIELD_SURVEY_LOG/ATTACHMENTS 메타에서 해당 src_id를 찾지 못했습니다: "
                        + ", ".join(missing)
                    ),
                    fix_hint="FIELD_SURVEY_LOG(조사일/조사자/범위 등) 또는 ATTACHMENTS(증빙 파일) 행에 동일 src_id를 입력해 근거를 연결",
                )
            )
        missing_sc = sorted(
            sid for sid in field_claim_cited_src_ids_in_summary_or_conclusion if sid not in linked_src_ids
        )
        if missing_sc:
            results.append(
                RuleResult(
                    rule_id="W-FIELD-META-002",
                    severity="WARN",
                    message=(
                        "결론/요약 섹션에서 현장 근거 출처를 인용했으나(case 본문), "
                        "case.xlsx의 FIELD_SURVEY_LOG/ATTACHMENTS 메타에서 해당 src_id를 찾지 못했습니다: "
                        + ", ".join(missing_sc)
                    ),
                    fix_hint="결론/요약 근거로 쓰인 현장 출처는 FIELD_SURVEY_LOG 또는 ATTACHMENTS 메타로 반드시 연결(제출 모드에서 ERROR로 승격)",
                )
            )

    # Source Register 정합성(본문 인용 → sources 등록부 존재)
    unknown: set[str] = set()
    citation_counter: Counter[str] = Counter()
    for s in qa_sections:
        for p in s.paragraphs:
            for cid in _extract_citation_ids(p):
                citation_counter[cid] += 1
                if cid in {"S-TBD", "SRC-TBD"}:
                    continue
                if not sources.has(cid):
                    unknown.add(cid)
    if unknown:
        results.append(
            RuleResult(
                rule_id="E-SRC-REG-001",
                severity="ERROR",
                message="sources.yaml에 없는 Source ID 참조: " + ", ".join(sorted(unknown)),
                fix_hint="sources.yaml에 해당 ID를 등록하거나 본문 인용 ID를 수정",
            )
        )

    # 임시 출처(S-TBD) 사용 경고(본문 인용 기준)
    if citation_counter.get("S-TBD", 0) + citation_counter.get("SRC-TBD", 0) > 0:
        results.append(
            RuleResult(
                rule_id="W-SRC-TBD-001",
                severity="WARN",
                message="본문에서 임시 출처(S-TBD/SRC-TBD)를 인용하고 있습니다.",
                fix_hint="sources.yaml에 실제 출처를 등록하고 본문 인용을 교체",
            )
        )

    # 스코핑 제외 사유
    for item in case.scoping_matrix:
        if item.scoping_class == ScopingClass.EXCLUDE and item.exclude_reason.is_empty():
            results.append(
                RuleResult(
                    rule_id="E-SCOPE-001",
                    severity="ERROR",
                    message=f"평가제외 항목 제외사유 누락: {item.item_id}",
                    fix_hint="scoping_matrix.exclude_reason 입력",
                )
            )

    # DIA 주민탐문 체크(실무지침 반영 - 간이 규칙):
    # - DRR_BASE_HAZARD에서 interview_done=Y인 항목이 있으면
    #   DRR_INTERVIEWS가 3건 이상이며, residence_years >= 10인 항목이 3건 이상인지 검사.
    disaster = (case.model_extra or {}).get("disaster")
    if isinstance(disaster, dict):
        hazard_history = disaster.get("hazard_history") or []
        interviews = disaster.get("interviews") or []

        interview_required = False
        if isinstance(hazard_history, list):
            for row in hazard_history:
                try:
                    if isinstance(row, dict) and str(row.get("interview_done", {}).get("t", "")).strip() == "Y":
                        interview_required = True
                        break
                except Exception:
                    continue

        if interview_required:
            count_total = len(interviews) if isinstance(interviews, list) else 0
            count_10y = 0
            if isinstance(interviews, list):
                for row in interviews:
                    try:
                        years = None
                        if isinstance(row, dict):
                            years = (row.get("residence_years") or {}).get("v")
                        if years is None:
                            continue
                        if float(years) >= 10:
                            count_10y += 1
                    except Exception:
                        continue
            if count_total < 3 or count_10y < 3:
                results.append(
                    RuleResult(
                        rule_id="W-DIA-INT-REQ-001",
                        severity="WARN",
                        message=(
                            "주민탐문(interviews) 권장 요건 미충족: "
                            f"총 {count_total}건, 10년 이상 {count_10y}건 (권장: 3건 이상)"
                        ),
                        fix_hint="DRR_INTERVIEWS에 익명 응답자 3명 이상(거주 10년 이상) 입력 권장",
                    )
                )

        # DIA 대상지역 설정(사업지/상류/하류/주변) 커버리지 체크(권장):
        parts = disaster.get("target_area_parts") or []
        if isinstance(parts, list) and parts:
            present: set[str] = set()
            for row in parts:
                try:
                    part = str((row.get("part") or {}).get("t", "")).strip().upper()
                except Exception:
                    part = ""
                if part:
                    present.add(part)
            required = {"PROJECT", "UPSTREAM", "DOWNSTREAM", "SURROUNDING"}
            missing = sorted(required - present)
            if missing:
                results.append(
                    RuleResult(
                        rule_id="W-DIA-TA-PARTS-REQ-001",
                        severity="WARN",
                        message=f"평가대상지역 4구분(사업지/상류/하류/주변) 일부 누락: {', '.join(missing)}",
                        fix_hint="DRR_TARGET_AREA_PARTS에 PROJECT/UPSTREAM/DOWNSTREAM/SURROUNDING 4행 입력 권장",
                    )
                )

    # 현장조사 '진실 게이트'(권장):
    # - 현장/탐문 결과 표현이 등장했는데 FIELD_SURVEY_LOG가 비어있으면 경고.
    has_field_log = bool((case.model_extra or {}).get("field_survey_log"))
    if not has_field_log:
        for s in qa_sections:
            for i, p in enumerate(s.paragraphs):
                if any(pat.search(p or "") for _, pat, _ in _FIELD_CLAIM_PATTERNS):
                    results.append(
                        RuleResult(
                            rule_id="W-FIELD-LOG-001",
                            severity="WARN",
                            message="현장/탐문 관련 표현이 있으나 FIELD_SURVEY_LOG가 비어있습니다.",
                            fix_hint="FIELD_SURVEY_LOG에 조사일/조사범위/조사자/사진폴더 등 최소 1건 입력 권장",
                        )
                    )
                    break
            else:
                continue
            break

    # 숫자 값에 출처 필수
    tbd_src_numeric_paths: list[str] = []
    for path, val, src in _walk(case):
        if val is None:
            continue
        try:
            # treat numeric-ish only
            float(val)
        except Exception:
            continue
        src_list = src if isinstance(src, list) else [src] if isinstance(src, str) else []
        if not [s for s in src_list if str(s).strip()]:
            results.append(
                RuleResult(
                    rule_id="E-SRC-001",
                    severity="ERROR",
                    path=path,
                    message=f"숫자 값에 출처 누락: {path}",
                    fix_hint="해당 수치의 src(source_ids) 입력",
                )
            )
            # keep going to report all, but this can be noisy
            continue

        if any(str(s).strip() in {"S-TBD", "SRC-TBD"} for s in src_list):
            if len(tbd_src_numeric_paths) < 20:
                tbd_src_numeric_paths.append(path)

    if tbd_src_numeric_paths:
        results.append(
            RuleResult(
                rule_id="W-SRC-TBD-002",
                severity="WARN",
                message="숫자 값이 임시 출처(S-TBD)로만 연결되어 있습니다(예: "
                + ", ".join(tbd_src_numeric_paths[:8])
                + (", ..." if len(tbd_src_numeric_paths) > 8 else "")
                + ").",
                fix_hint="해당 수치의 실제 출처(src_id)를 sources.yaml에 등록하고 연결",
            )
        )

    # 표 스펙 기반 검증
    if spec is not None:
        tables = list(spec.tables.tables)
        if present_table_ids is not None:
            tables = [t for t in tables if t.id in present_table_ids]

        for t in tables:
            for finding in validate_table(case, t):
                results.append(
                    RuleResult(
                        rule_id=finding.rule_id,
                        severity=finding.severity,
                        message=finding.message,
                    )
                )

        # 표 입력 누락(빈 셀 placeholder) 탐지
        empty_cell_token = str(getattr(spec.tables.defaults, "empty_cell", "") or "").strip()
        if empty_cell_token:
            for t in tables:
                try:
                    td = build_table(case, sources, t, spec.tables.defaults)
                except Exception:
                    continue
                ph = 0
                for row in td.rows:
                    for cell in row:
                        cell_text = "" if cell is None else str(cell)
                        if cell_text.strip() == empty_cell_token:
                            ph += 1
                if ph:
                    results.append(
                        RuleResult(
                            rule_id="W-TBL-PLACEHOLDER-001",
                            severity="WARN",
                            path=t.id,
                            message=f"표 입력 누락 표시({empty_cell_token})가 남아있습니다(표 {t.id}, {ph}칸).",
                            fix_hint="case.xlsx(v2) 입력값을 채우거나, 해당 표를 REUSE/AUTO로 전환",
                        )
                    )

        # DIA strictness toggle (for self-use vs. submission readiness):
        # When PROJECT.doc_drr_required == "Y", enforce core DIA numeric inputs even when
        # table columns are allowed to be empty (clean draft output).
        report_type = str(getattr(getattr(spec.sections, "doc_profile", None), "report_type", "") or "").strip()
        doc_drr_required = str(getattr(case.meta, "doc_drr_required", "") or "").strip()
        if ("재해" in report_type) and doc_drr_required.upper() == "Y":
            rainfall = resolve_path(case, "disaster.rainfall")
            runoff = resolve_path(case, "disaster.runoff_basins")
            sediment = resolve_path(case, "disaster.sediment_erosion")

            rainfall_rows = rainfall if isinstance(rainfall, list) else []
            runoff_rows = runoff if isinstance(runoff, list) else []
            sediment_rows = sediment if isinstance(sediment, list) else []

            has_return_period = any(resolve_path(r, "frequency_year.v") is not None for r in rainfall_rows)
            if rainfall_rows and not has_return_period:
                results.append(
                    RuleResult(
                        rule_id="E-DIA-STRICT-RAIN-001",
                        severity="ERROR",
                        path="DRR_HYDRO_RAIN.return_period_yr",
                        message="DIA required(Y)인데 강우 빈도(재현기간)가 비어 있습니다.",
                        fix_hint="case.xlsx(v2) `DRR_HYDRO_RAIN.return_period_yr`를 채우고(근거 포함) 재실행",
                    )
                )

            missing_cn = 0
            missing_pre = 0
            missing_post = 0
            for r in runoff_rows:
                if resolve_path(r, "cn_value.v") is None:
                    missing_cn += 1
                if resolve_path(r, "pre_peak_cms.v") is None:
                    missing_pre += 1
                if resolve_path(r, "post_peak_cms.v") is None:
                    missing_post += 1
            if runoff_rows and (missing_cn or missing_pre or missing_post):
                results.append(
                    RuleResult(
                        rule_id="E-DIA-STRICT-RUNOFF-001",
                        severity="ERROR",
                        path="DRR_HYDRO_MODEL",
                        message=(
                            "DIA required(Y)인데 유출 해석 핵심 수치가 비어 있습니다("
                            f"CN 누락 {missing_cn}건, 첨두유출(전) 누락 {missing_pre}건, 첨두유출(후) 누락 {missing_post}건)."
                        ),
                        fix_hint="case.xlsx(v2) `DRR_HYDRO_MODEL.cn_or_c/peak_cms_before/peak_cms_after`를 계산서 근거로 입력",
                    )
                )

            has_soil_loss = any(
                (resolve_path(r, "soil_loss_before.v") is not None) or (resolve_path(r, "soil_loss_after.v") is not None)
                for r in sediment_rows
            )
            if sediment_rows and not has_soil_loss:
                results.append(
                    RuleResult(
                        rule_id="E-DIA-STRICT-SED-001",
                        severity="ERROR",
                        path="DRR_SEDIMENT",
                        message="DIA required(Y)인데 토사유출/침식 산정 결과(ton/ha/yr)가 비어 있습니다.",
                        fix_hint="case.xlsx(v2) `DRR_SEDIMENT.soil_loss_t_ha_yr_before/after`를 산정 근거로 입력",
                    )
                )

        # 그림(필수/조건부) 검증
        figures = list(spec.figures.figures)
        if present_figure_ids is not None:
            figures = [f for f in figures if f.id in present_figure_ids]

        fig_map = build_figure_map(case, figures, asset_search_dirs=asset_search_dirs)

        # Reference guardrail (best-effort): enforce caption hint when source_origin=REFERENCE.
        for f in figures:
            resolved = fig_map.get(f.id)
            if not resolved:
                continue
            gm = str(getattr(resolved, "gen_method", "") or "").strip()
            is_ref_by_origin = _is_reference_figure(case, f.id)
            is_ref_by_gen_method = bool(gm and _AUTH_REFERENCE_RE.search(gm))
            if not (is_ref_by_origin or is_ref_by_gen_method):
                continue

            # If the user explicitly forced AUTHENTICITY:OFFICIAL while marking REFERENCE,
            # the watermark/caption guardrail may not apply. Catch this mismatch.
            if is_ref_by_origin and not is_ref_by_gen_method:
                results.append(
                    RuleResult(
                        rule_id="E-FIG-REF-002",
                        severity="ERROR",
                        path=f.id,
                        message=f"REFERENCE 그림인데 워터마크 힌트(AUTHENTICITY:REFERENCE)가 누락되었습니다: {f.id}",
                        fix_hint="FIGURES.source_origin=REFERENCE면 FIGURES.gen_method에 AUTHENTICITY:OFFICIAL을 넣지 말고, 비워두거나 AUTHENTICITY:REFERENCE로 정합성을 맞추세요.",
                    )
                )

            if is_required(case, f):
                results.append(
                    RuleResult(
                        rule_id="W-FIG-REF-003",
                        severity="WARN",
                        path=f.id,
                        message=f"필수 그림이 REFERENCE로 표시되어 있습니다(공식 도면 필요 가능성): {f.id}",
                        fix_hint="가능하면 공식 도면/공식 DB 기반 자료로 교체하고, 참고도는 워터마크/캡션과 함께 DISPLAY_ONLY 용도로만 사용하세요.",
                    )
                )

            cap = str(getattr(resolved, "caption", "") or "").strip()
            if not cap:
                continue
            if any(token in cap for token in ["참고", "REFERENCE", "공식 도면 아님"]):
                continue
            results.append(
                RuleResult(
                    rule_id="W-FIG-REF-001",
                    severity="WARN",
                    path=f.id,
                    message=f"REFERENCE 그림 캡션에 참고도 표기가 없습니다: {f.id}",
                    fix_hint="FIGURES.source_origin=REFERENCE면 캡션에 '참고도' 표기 권장(가드레일).",
                )
            )

        for f in figures:
            if not is_required(case, f):
                continue
            resolved = fig_map.get(f.id)
            if not resolved or not resolved.file_path:
                results.append(
                    RuleResult(
                        rule_id="E-FIG-001",
                        severity="ERROR",
                        message=f"필수 그림 누락: {f.id} ({f.caption})",
                        fix_hint=f"assets에 type={f.asset_type} 파일 등록",
                    )
                )

        # PDF page auto-selection: when no explicit page is provided, we pick a likely drawing page.
        # Expose the selected page as a WARN so users can lock it down (page drift risk).
        for f in figures:
            resolved = fig_map.get(f.id)
            if not resolved or not resolved.file_path:
                continue
            try:
                p = Path(str(resolved.file_path))
            except Exception:
                continue
            if p.suffix.lower() != ".pdf":
                continue

            try:
                page, page_source = select_pdf_page(
                    p, gen_method=getattr(resolved, "gen_method", None), target_dpi=120, max_pages=15
                )
            except Exception as e:
                results.append(
                    RuleResult(
                        rule_id="W-FIG-PDF-001",
                        severity="WARN",
                        path=f.id,
                        message=f"PDF 페이지 자동선택(휴리스틱) 실패: {f.id} ({type(e).__name__})",
                        fix_hint="PDF 렌더링 의존성(Python fitz/PyMuPDF)을 설치하거나, 해당 그림을 이미지(PNG/JPG)로 제공",
                    )
                )
                continue

            if page_source != "explicit":
                results.append(
                    RuleResult(
                        rule_id="W-FIG-PDF-001",
                        severity="WARN",
                        path=f.id,
                        message=f"PDF 페이지 미지정: {f.id} → 자동 선택 page={page} (max_pages=15)",
                        fix_hint=f"FIGURES.file_path에 '#page={page}'를 추가하거나 FIGURES.gen_method='PDF_PAGE:{page}'로 고정",
                    )
                )

            # A3 guideline (best-effort): warn when the selected PDF page is larger than A3.
            size_mm = _pdf_page_size_mm(p, page_1based=page)
            if size_mm:
                w_mm, h_mm = size_mm
                long_mm = max(w_mm, h_mm)
                short_mm = min(w_mm, h_mm)
                if long_mm > _A3_LONG_MM + 5.0 or short_mm > _A3_SHORT_MM + 5.0:
                    results.append(
                        RuleResult(
                            rule_id="W-FIG-PDF-SIZE-001",
                            severity="WARN",
                            path=f.id,
                            message=f"도면 PDF 권장규격(A3 이하) 초과 가능성: {f.id} (page={page}, {w_mm:.0f}×{h_mm:.0f}mm)",
                            fix_hint="가능하면 A3 이하 도면으로 분할/축소한 PDF를 제공(가독성/제출 일관성).",
                        )
                    )

    # “비었을 때 어디서 가져오나” 규칙(차기; config-driven).
    # - Evaluates `config/data_acquisition_rules.yaml` against the input case.xlsx(v2)
    # - Emits RuleResult rows so they appear in validation_report*.json and VALIDATION_SUMMARY.
    if case_xlsx_path is not None:
        try:
            xlsx_path = Path(case_xlsx_path).resolve()
        except Exception:
            xlsx_path = None

        rules_path = data_acquisition_rules_path or _default_data_acquisition_rules_path()
        if xlsx_path and xlsx_path.exists() and rules_path and rules_path.exists():
            rules = _load_data_acquisition_rules(rules_path)
            if rules:
                results.extend(
                    _evaluate_data_acquisition_rules(
                        xlsx_path=xlsx_path, rules=rules, submission_mode=submission_mode
                    )
                )

    # Submission mode: fail fast on any remaining placeholders/TODOs/temporary sources.
    # This intentionally upgrades some WARNs to ERRORs so a single `--submission` run can be
    # interpreted as "ready to submit" (no TODO, no placeholders, no S-TBD citations).
    if submission_mode:
        strict_rule_ids = {
            "W-TODO-001",
            "W-PLACEHOLDER-001",
            "W-TBL-PLACEHOLDER-001",
            "W-SRC-TBD-001",
            "W-SRC-TBD-002",
            "W-FIELD-META-002",
        }
        for r in results:
            if r.severity == "WARN" and r.rule_id in strict_rule_ids:
                r.severity = "ERROR"  # type: ignore[assignment]

    # 통계
    stats: dict[str, int] = {}
    stats["placeholder_count"] = all_text.count("【작성자 기입 필요】") + all_text.count("[작성자 기입 필요]")
    stats["s_tbd_citation_count"] = citation_counter.get("S-TBD", 0) + citation_counter.get("SRC-TBD", 0)
    stats["error_count"] = len([r for r in results if r.severity == "ERROR"])
    stats["warn_count"] = len([r for r in results if r.severity == "WARN"])

    return ValidationReport(results=results, stats=stats)
