from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches


def pdf_to_image_docx(*, pdf_path: Path, out_docx_path: Path, dpi: int = 200) -> None:
    """Create a DOCX that visually matches the PDF by embedding each page as a full-page image.

    Notes:
    - Output is pixel-perfect (layout), but the text is not editable/searchable (it's an image).
    - Intended as a strict "format identical" option when editable DOCX is not required.
    """
    import fitz  # PyMuPDF

    pdf_path = pdf_path.expanduser().resolve()
    out_docx_path = out_docx_path.expanduser().resolve()
    out_docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    with fitz.open(str(pdf_path)) as pdf:
        if pdf.page_count == 0:
            raise ValueError(f"PDF has no pages: {pdf_path}")

        # Assume consistent page size; set from page 1.
        first = pdf.load_page(0)
        w_in = float(first.rect.width) / 72.0
        h_in = float(first.rect.height) / 72.0

        section = doc.sections[0]
        section.page_width = Inches(w_in)
        section.page_height = Inches(h_in)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        for i in range(pdf.page_count):
            page = pdf.load_page(i)
            pix = page.get_pixmap(dpi=int(dpi), alpha=False)
            img = BytesIO(pix.tobytes("png"))

            p = doc.add_paragraph()
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
            run = p.add_run()
            run.add_picture(img, width=section.page_width)

            if i != pdf.page_count - 1:
                doc.add_page_break()

    doc.save(str(out_docx_path))

