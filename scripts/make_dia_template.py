from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="templates/dia_template.docx")
    args = parser.parse_args()

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("소규모재해영향평가서(재해영향성검토서) 템플릿", level=1)
    doc.add_paragraph("※ 본 문서는 앵커 치환용 템플릿입니다(SSOT: spec_dia/*.yaml).")
    doc.add_page_break()

    doc.add_heading("표지", level=1)
    doc.add_paragraph("[[BLOCK:DIA0_COVER]]")

    doc.add_heading("요약", level=1)
    doc.add_paragraph("[[BLOCK:DIA0_SUMMARY]]")

    doc.add_heading("제1장 사업의 개요", level=1)
    doc.add_paragraph("[[BLOCK:DIA1_PROJECT]]")
    doc.add_paragraph("[[TABLE:PARCELS]]")
    doc.add_paragraph("[[TABLE:FACILITIES]]")
    doc.add_paragraph("[[TABLE:SCHEDULE]]")

    doc.add_heading("제2장 평가대상지역 설정", level=1)
    doc.add_paragraph("[[BLOCK:DIA2_TARGET_AREA]]")
    doc.add_paragraph("[[TABLE:DIA_TARGET_AREA]]")
    doc.add_paragraph("[[TABLE:DIA_TARGET_AREA_PARTS]]")
    doc.add_paragraph("[[TABLE:DIA_SCOPE]]")
    doc.add_paragraph("[[FIG:FIG-DIA-TARGET-01]]")

    doc.add_heading("제3장 재해 관련 기초현황", level=1)
    doc.add_paragraph("[[BLOCK:DIA3_BASELINE]]")
    doc.add_paragraph("[[TABLE:DIA_RAINFALL]]")
    doc.add_paragraph("[[TABLE:DIA_BASE_HAZARD]]")
    doc.add_paragraph("[[TABLE:DIA_INTERVIEWS]]")
    doc.add_paragraph("[[TABLE:DIA_DRAINAGE]]")
    doc.add_paragraph("[[FIG:FIG-DIA-DRAINAGE-01]]")
    doc.add_paragraph("[[FIG:FIG-DIA-STORMWATER-01]]")

    doc.add_heading("제4장 재해영향 분석", level=1)
    doc.add_paragraph("[[BLOCK:DIA4_ANALYSIS]]")
    doc.add_paragraph("[[TABLE:DIA_RUNOFF]]")
    doc.add_paragraph("[[TABLE:DIA_SEDIMENT]]")
    doc.add_paragraph("[[TABLE:DIA_SLOPE]]")

    doc.add_heading("제5장 재해 저감대책", level=1)
    doc.add_paragraph("[[BLOCK:DIA5_MITIGATION]]")
    doc.add_paragraph("[[TABLE:DIA_MITIGATION]]")

    doc.add_heading("제6장 유지관리계획 및 유지관리대장", level=1)
    doc.add_paragraph("[[BLOCK:DIA6_MAINTENANCE]]")
    doc.add_paragraph("[[TABLE:DIA_MAINTENANCE]]")

    doc.add_heading("제7장 종합결론", level=1)
    doc.add_paragraph("[[BLOCK:DIA7_CONCLUSION]]")

    doc.add_heading("부록", level=1)
    doc.add_paragraph("[[TABLE:SOURCE_REGISTER]]")

    doc.save(out)
    print(f"wrote: {out}")


if __name__ == "__main__":
    main()
