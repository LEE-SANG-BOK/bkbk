from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Mm


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="templates/report_template.docx")
    args = parser.parse_args()

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    doc.add_heading("소규모환경영향평가서(관광농원) 템플릿", level=1)
    doc.add_paragraph("※ 본 문서는 앵커 치환용 템플릿입니다(SSOT: spec/*.yaml).")
    doc.add_page_break()

    doc.add_heading("표지", level=1)
    doc.add_paragraph("[[BLOCK:CH0_COVER]]")

    doc.add_heading("요약", level=1)
    doc.add_paragraph("[[BLOCK:CH0_SUMMARY]]")

    doc.add_heading("제1장 사업의 개요", level=1)
    doc.add_heading("1. 사업의 목적 및 필요성", level=2)
    doc.add_paragraph("[[BLOCK:CH1_PURPOSE]]")
    doc.add_heading("2. 사업의 위치 및 면적", level=2)
    doc.add_paragraph("[[BLOCK:CH1_LOCATION_AREA]]")
    doc.add_paragraph("[[TABLE:PARCELS]]")
    doc.add_paragraph("[[TABLE:FACILITIES]]")
    doc.add_paragraph("[[TABLE:ZONING_BREAKDOWN]]")
    doc.add_paragraph("[[FIG:FIG-LOC-01]]")
    doc.add_paragraph("[[FIG:FIG-LAYOUT-01]]")

    doc.add_heading("3. 사업 내용 및 규모", level=2)
    doc.add_paragraph("[[BLOCK:CH1_SCALE]]")

    doc.add_heading("4. 사업 추진 일정", level=2)
    doc.add_paragraph("[[BLOCK:CH1_SCHEDULE]]")
    doc.add_paragraph("[[TABLE:SCHEDULE]]")

    doc.add_heading("5. 인허가 현황 및 협의 대상 근거", level=2)
    doc.add_paragraph("[[BLOCK:CH1_APPLICABILITY]]")

    doc.add_heading("제2장 환경현황 조사", level=1)
    doc.add_heading("1. 조사 범위·방법", level=2)
    doc.add_paragraph("[[BLOCK:CH2_METHOD]]")
    doc.add_paragraph("[[TABLE:ZONING_OVERLAY]]")
    doc.add_paragraph("[[FIG:FIG-IA-01]]")

    doc.add_heading("환경기준(기준표)", level=2)
    doc.add_paragraph("[[TABLE:ENV_AIR_STANDARDS]]")
    doc.add_paragraph("[[TABLE:ENV_LIVING_STANDARDS]]")
    doc.add_paragraph("[[TABLE:ENV_NOISE_STANDARDS]]")

    doc.add_heading("2. 자연환경(지형·지질)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_TOPO]]")
    doc.add_heading("3. 자연환경(동·식물상)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_ECO]]")
    doc.add_paragraph("[[FIG:FIG-ECO-ROUTE-01]]")
    doc.add_paragraph("[[FIG:FIG-ECO-PHOTO-01]]")
    doc.add_heading("4. 자연환경(수환경)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_WATER]]")

    doc.add_heading("5. 생활환경(대기질)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_AIR]]")
    doc.add_heading("6. 생활환경(소음·진동)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_NOISE]]")
    doc.add_heading("7. 생활환경(악취)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_ODOR]]")

    doc.add_heading("8. 사회·경제(토지이용)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_LANDUSE]]")
    doc.add_paragraph("[[FIG:FIG-LANDUSE-01]]")
    doc.add_paragraph("[[FIG:FIG-AERIAL-01]]")
    doc.add_heading("9. 사회·경제(경관)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_LANDSCAPE]]")
    doc.add_paragraph("[[FIG:FIG-VP-01]]")
    doc.add_paragraph("[[FIG:FIG-VP-02]]")
    doc.add_heading("10. 사회·경제(인구·주거/교통)", level=2)
    doc.add_paragraph("[[BLOCK:CH2_POP_TRAFFIC]]")

    doc.add_heading("환경현황 요약표", level=2)
    doc.add_paragraph("[[TABLE:BASELINE_SUMMARY]]")

    doc.add_heading("제3장 환경영향 예측", level=1)
    doc.add_heading("1. 평가항목 선정", level=2)
    doc.add_paragraph("[[BLOCK:CH3_SCOPING]]")
    doc.add_paragraph("[[TABLE:SCOPING]]")
    doc.add_heading("2. 공사 단계", level=2)
    doc.add_paragraph("[[BLOCK:CH3_CONSTRUCTION]]")
    doc.add_heading("3. 운영 단계", level=2)
    doc.add_paragraph("[[BLOCK:CH3_OPERATION]]")

    doc.add_heading("제4장 환경보전 및 저감방안", level=1)
    doc.add_paragraph("[[BLOCK:CH4_MITIGATION]]")
    doc.add_paragraph("[[FIG:FIG-DRAINAGE-01]]")
    doc.add_paragraph("[[TABLE:MITIGATION_PLAN]]")

    doc.add_heading("제5장 환경관리 계획(협의조건 이행관리)", level=1)
    doc.add_paragraph("[[BLOCK:CH5_TRACKER]]")
    doc.add_paragraph("[[TABLE:CONDITION_TRACKER]]")
    doc.add_paragraph("[[BLOCK:CH5_2_ORG]]")
    doc.add_paragraph("[[BLOCK:CH5_3_MONITORING]]")

    doc.add_heading("제6장 주민의견 수렴 결과(해당 시)", level=1)
    doc.add_paragraph("[[BLOCK:CH6_PUBLIC]]")

    doc.add_heading("제7장 종합평가 및 결론", level=1)
    doc.add_paragraph("[[BLOCK:CH7_CONCLUSION]]")

    doc.add_heading("부록", level=1)
    doc.add_paragraph("[[TABLE:SOURCE_REGISTER]]")

    doc.save(out)
    print(f"wrote: {out}")


if __name__ == "__main__":
    main()
