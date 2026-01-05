from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


def _set_default_styles(doc: Document) -> None:
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    # Caption (used for <표 ...>, <그림 ...>)
    if "Caption" in styles:
        cap = styles["Caption"]
    else:
        cap = styles.add_style("Caption", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
    cap.font.name = "맑은 고딕"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    cap.font.size = Pt(10)


def _set_page_margins(doc: Document) -> None:
    sec = doc.sections[0]
    # Match Korean submission norm + approved sample PDF(A4).
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)


def _add_anchor_line(doc: Document, anchor: str, align_center: bool = False) -> None:
    p = doc.add_paragraph(anchor)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="templates/report_template.sample_changwon_2025.docx")
    args = ap.parse_args()

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_default_styles(doc)
    _set_page_margins(doc)

    # Front matter (sample-like: title + summary + TOC placeholder)
    doc.add_heading("소규모환경영향평가서(관광농원)", level=1)
    doc.add_paragraph("※ 본 문서는 ‘샘플(창원, 2025)’ 서식에 맞춘 앵커 치환용 템플릿입니다.")
    doc.add_paragraph("※ 내용은 case.xlsx + sources.yaml + attachments 기반으로 자동 작성됩니다.")
    doc.add_page_break()

    doc.add_heading("표지", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH0_COVER]]")
    doc.add_page_break()

    doc.add_heading("요약", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH0_SUMMARY]]")
    doc.add_page_break()

    doc.add_heading("목차", level=1)
    doc.add_paragraph("※ Word 목차(자동) 사용 시, 생성 후 Word에서 ‘목차 업데이트’ 수행")
    doc.add_page_break()

    # Chapter 1 (사업의 개요)
    doc.add_heading("제1장 사업의 개요", level=1)
    doc.add_heading("1. 사업의 목적 및 필요성", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH1_PURPOSE]]")
    doc.add_heading("2. 사업의 위치 및 면적", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH1_LOCATION_AREA]]")
    _add_anchor_line(doc, "[[TABLE:PARCELS]]")
    _add_anchor_line(doc, "[[TABLE:ZONING_BREAKDOWN]]")
    _add_anchor_line(doc, "[[FIG:FIG-LOC-01]]", align_center=True)
    doc.add_heading("3. 사업 내용 및 규모", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH1_SCALE]]")
    _add_anchor_line(doc, "[[TABLE:FACILITIES]]")
    _add_anchor_line(doc, "[[FIG:FIG-LAYOUT-01]]", align_center=True)
    doc.add_heading("4. 사업 추진 일정", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH1_SCHEDULE]]")
    _add_anchor_line(doc, "[[TABLE:SCHEDULE]]")
    doc.add_heading("5. 인허가 현황 및 협의 대상 근거", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH1_APPLICABILITY]]")
    doc.add_page_break()

    # Chapter 2 (지역개황/기준/규제 등 — sample has rich tables here)
    doc.add_heading("제2장 지역개황", level=1)
    doc.add_heading("1. 조사 범위·방법", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_METHOD]]")
    _add_anchor_line(doc, "[[FIG:FIG-IA-01]]", align_center=True)
    doc.add_heading("2. 환경관련 지구·지역(보호/규제) 현황", level=2)
    doc.add_paragraph("※ ‘해당(O/X) + 이격거리’ 총괄표는 샘플 서식에 맞춰 자동 생성됩니다.")
    _add_anchor_line(doc, "[[TABLE:ZONING_OVERLAY]]")
    doc.add_heading("3. 환경기준 및 보호대상시설(요약)", level=2)
    doc.add_paragraph("※ 기준/보호대상시설은 공공자료/지침 기반으로 작성하며 출처를 함께 표기합니다.")
    _add_anchor_line(doc, "[[TABLE:ENV_AIR_STANDARDS]]")
    _add_anchor_line(doc, "[[TABLE:ENV_LIVING_STANDARDS]]")
    _add_anchor_line(doc, "[[TABLE:ENV_NOISE_STANDARDS]]")
    doc.add_page_break()

    # Chapter 3 (대상사업의 지역 범위 + 평가항목 선정)
    doc.add_heading("제3장 대상사업의 지역 범위", level=1)
    doc.add_heading("1. 평가항목 선정(중점/현황/제외) 및 사유", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH3_SCOPING]]")
    _add_anchor_line(doc, "[[TABLE:SCOPING]]")
    doc.add_page_break()

    # Chapter 4 (주변지역 토지이용)
    doc.add_heading("제4장 대상지역의 주변지역에 대한 토지이용 현황", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH2_LANDUSE]]")
    _add_anchor_line(doc, "[[FIG:FIG-LANDUSE-01]]", align_center=True)
    _add_anchor_line(doc, "[[FIG:FIG-AERIAL-01]]", align_center=True)
    doc.add_page_break()

    # Chapter 5 (환경 현황)
    doc.add_heading("제5장 환경 현황", level=1)
    doc.add_heading("1. 자연환경(지형·지질)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_TOPO]]")
    doc.add_heading("2. 자연환경(동·식물상)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_ECO]]")
    _add_anchor_line(doc, "[[FIG:FIG-ECO-ROUTE-01]]", align_center=True)
    _add_anchor_line(doc, "[[FIG:FIG-ECO-PHOTO-01]]", align_center=True)
    doc.add_heading("3. 자연환경(수환경)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_WATER]]")
    doc.add_heading("4. 생활환경(대기질)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_AIR]]")
    doc.add_heading("5. 생활환경(소음·진동)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_NOISE]]")
    doc.add_heading("6. 생활환경(악취)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_ODOR]]")
    doc.add_heading("7. 사회·경제(경관)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_LANDSCAPE]]")
    _add_anchor_line(doc, "[[FIG:FIG-VP-01]]", align_center=True)
    _add_anchor_line(doc, "[[FIG:FIG-VP-02]]", align_center=True)
    doc.add_heading("8. 사회·경제(인구·주거/교통)", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH2_POP_TRAFFIC]]")
    doc.add_heading("9. 환경현황 요약표", level=2)
    _add_anchor_line(doc, "[[TABLE:BASELINE_SUMMARY]]")
    doc.add_page_break()

    # Chapter 6 (입지의 타당성)
    doc.add_heading("제6장 입지의 타당성", level=1)
    doc.add_paragraph("※ 본 장은 입력된 규제/보호지역/민감수용체/토지이용계획 등을 근거로 요약합니다.")
    doc.add_paragraph("【작성자 기입 필요】(v0.1에서는 결론/근거만 자동 요약)")
    doc.add_page_break()

    # Chapter 7 (영향예측 및 환경보전방안)
    doc.add_heading("제7장 환경 현황과 환경에 미치는 영향의 조사·예측·평가 및 환경보전방안", level=1)
    doc.add_heading("1. 공사 단계", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH3_CONSTRUCTION]]")
    doc.add_heading("2. 운영 단계", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH3_OPERATION]]")
    doc.add_heading("3. 환경보전 및 저감방안", level=2)
    _add_anchor_line(doc, "[[BLOCK:CH4_MITIGATION]]")
    _add_anchor_line(doc, "[[FIG:FIG-DRAINAGE-01]]", align_center=True)
    _add_anchor_line(doc, "[[TABLE:MITIGATION_PLAN]]")
    doc.add_page_break()

    # Chapter 8 (환경관리 계획)
    doc.add_heading("제8장 협의조건 이행관리 및 환경관리 계획", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH5_TRACKER]]")
    _add_anchor_line(doc, "[[TABLE:CONDITION_TRACKER]]")
    _add_anchor_line(doc, "[[BLOCK:CH5_2_ORG]]")
    _add_anchor_line(doc, "[[BLOCK:CH5_3_MONITORING]]")
    doc.add_page_break()

    # Chapter 9 (주민의견 수렴)
    doc.add_heading("제9장 주민의견 수렴 결과(해당 시)", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH6_PUBLIC]]")
    doc.add_page_break()

    # Chapter 10 (결론/부록)
    doc.add_heading("제10장 종합평가 및 결론", level=1)
    _add_anchor_line(doc, "[[BLOCK:CH7_CONCLUSION]]")
    doc.add_page_break()

    doc.add_heading("부록", level=1)
    _add_anchor_line(doc, "[[BLOCK:APPENDIX_INSERTS]]")
    doc.add_heading("출처/근거 관리표(Source Register)", level=2)
    _add_anchor_line(doc, "[[TABLE:SOURCE_REGISTER]]")

    doc.save(out)
    print(f"wrote: {out}")


if __name__ == "__main__":
    main()
